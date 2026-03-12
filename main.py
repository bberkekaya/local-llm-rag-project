from __future__ import annotations

import io
import json
import logging
import os
import uuid
from functools import lru_cache
from pathlib import Path
from typing import List

import chromadb
import requests
from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from haystack import Document
from haystack.components.builders import PromptBuilder
from haystack_integrations.components.embedders.ollama import OllamaTextEmbedder
from haystack_integrations.components.generators.ollama import OllamaGenerator
from pydantic import BaseModel, Field
from pypdf import PdfReader
from docx import Document as DocxDocument


logger = logging.getLogger(__name__)
if not logger.handlers:
    logging.basicConfig(level=logging.INFO)


# ---- Model seçimleri ----
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "bge-m3:567m")
LLM_MODEL = os.getenv("LLM_MODEL", "qwen3.5:9b")
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
GENERATION_NUM_PREDICT = int(os.getenv("GENERATION_NUM_PREDICT", "1024"))
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1200"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))
MAX_FILE_SIZE_MB = 50
MAX_FILE_COUNT = 10
MAX_QUERY_LENGTH = 2000
MAX_SOURCE_LENGTH = 180
ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx", ".doc"}
ALLOWED_CONTENT_TYPES = {
    "text/plain",
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
DEFAULT_ALLOWED_ORIGINS = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
    "http://localhost:8000",
    "http://127.0.0.1:8000",
]

# ---- ChromaDB konumu ----
PERSIST_DIR = Path("chromadb_store")
PERSIST_DIR.mkdir(exist_ok=True)



_chroma_client = None


def get_allowed_origins() -> list[str]:
    configured = os.getenv("ALLOWED_ORIGINS", "")
    if configured.strip():
        return [origin.strip() for origin in configured.split(",") if origin.strip()]
    return DEFAULT_ALLOWED_ORIGINS


def sanitize_filename(filename: str | None) -> str:
    if not filename:
        raise HTTPException(status_code=400, detail="Geçersiz dosya adı.")

    name = Path(filename).name.strip().replace("\x00", "")
    if not name or name in {".", ".."}:
        raise HTTPException(status_code=400, detail="Geçersiz dosya adı.")
    if len(name) > MAX_SOURCE_LENGTH:
        raise HTTPException(status_code=400, detail="Dosya adı çok uzun.")
    return name


def build_safe_error(message: str, status_code: int = 500) -> HTTPException:
    return HTTPException(status_code=status_code, detail=message)


def build_ollama_model_error(model_name: str) -> HTTPException:
    return HTTPException(
        status_code=503,
        detail=(
            f'Ollama modeli bulunamadı: {model_name}. '
            f"Modeli önce host makinede `ollama pull {model_name}` komutuyla indirin."
        ),
    )


def normalize_ollama_exception(exc: Exception, model_name: str) -> HTTPException:
    message = str(exc).lower()
    if "not found" in message or "pull it first" in message:
        return build_ollama_model_error(model_name)
    return build_safe_error("Ollama servisi işlenemeyen bir hata döndürdü.", status_code=502)

def get_chroma_client():
    global _chroma_client
    if _chroma_client is None:
        _chroma_client = chromadb.PersistentClient(path=str(PERSIST_DIR))
    return _chroma_client


def get_collection():
    client = get_chroma_client()
    # bge-m3:567m modeli 1024 boyutlu embedding üretir
    return client.get_or_create_collection(
        name="documents",
        metadata={"hnsw:space": "cosine"}
    )


@lru_cache(maxsize=1)
def load_embedder() -> OllamaTextEmbedder:
    embedder = OllamaTextEmbedder(
        model=EMBEDDING_MODEL,
        url=OLLAMA_URL,
    )
    return embedder


@lru_cache(maxsize=1)
def load_generator() -> OllamaGenerator:
    generator = OllamaGenerator(
        model=LLM_MODEL,
        url=OLLAMA_URL,
        generation_kwargs={
            "num_predict": GENERATION_NUM_PREDICT,
            "temperature": 0.7,  # Yaratıcılık seviyesi
        },
    )
    return generator


def read_txt(file: io.BytesIO) -> str:
    return file.read().decode("utf-8", errors="ignore")


def read_pdf(file: io.BytesIO) -> str:
    reader = PdfReader(file)
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def read_docx(file: io.BytesIO) -> str:
    doc = DocxDocument(file)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    return "\n".join(paragraphs)


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    if chunk_size <= overlap:
        raise ValueError("chunk_size overlap değerinden büyük olmalıdır.")

    chunks = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def embed_and_store(documents: List[Document], collection, embedder) -> int:
    added = 0
    for doc in documents:
        source = sanitize_filename(doc.meta.get("source", "upload"))
        chunks = list(dict.fromkeys(chunk_text(doc.content)))
        if not chunks:
            continue

        existing = collection.get(where={"source": source}, include=[])
        existing_ids = existing.get("ids", [])
        if existing_ids:
            collection.delete(ids=existing_ids)

        ids = []
        embeddings = []
        metadatas = []
        for chunk in chunks:
            embedding = embedder.run(text=chunk)["embedding"]
            ids.append(str(uuid.uuid4()))
            embeddings.append(embedding)
            metadatas.append({"source": source})

        collection.add(
            ids=ids,
            documents=chunks,
            embeddings=embeddings,
            metadatas=metadatas,
        )
        added += len(chunks)
    return added


def build_prompt(query: str, contexts: List[str], style_instruction: str) -> str:
    template = """Aşağıdaki bağlamlara göre soruyu yanıtla.

Bağlamlar:
{% for doc in documents %}
- {{ doc }}
{% endfor %}

Soru: {{ query }}

Cevap (Türkçe):
- Cevap stili: {{ style_instruction }}"""
    prompt_builder = PromptBuilder(template=template)
    return prompt_builder.run(query=query, documents=contexts, style_instruction=style_instruction)["prompt"]


class QueryRequest(BaseModel):
    query: str = Field(min_length=1, max_length=MAX_QUERY_LENGTH)
    top_k: int = Field(default=4, ge=1, le=10)
    answer_style: str | None = Field(default=None, max_length=20)


app = FastAPI(title="RAG API (FastAPI + Chroma + Haystack)", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=get_allowed_origins(),
    allow_credentials=False,
    allow_methods=["GET", "POST", "DELETE"],
    allow_headers=["Accept", "Content-Type"],
)


@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "same-origin"
    response.headers["Cache-Control"] = "no-store"
    return response

@app.get("/")
def root():
    return {
        "message": "RAG API çalışıyor. Sağlık kontrolü: /api/health, embed: /api/embed, sorgu: /api/query"
    }


@app.get("/api/health")
def health():
    return {"status": "ok", "embedding_model": EMBEDDING_MODEL, "llm_model": LLM_MODEL}


@app.get("/api/files")
def list_files():
    try:
        collection = get_collection()
        data = collection.get(include=["metadatas"], limit=100000)
        sources = {}
        for meta in data.get("metadatas", []):
            if not meta:
                continue
            src = meta.get("source")
            if not src:
                continue
            sources[src] = sources.get(src, 0) + 1
        items = [{"source": k, "chunks": v} for k, v in sorted(sources.items())]
        return {"files": items, "total_chunks": sum(sources.values())}
    except Exception:
        logger.exception("Dosya listeleme hatası")
        raise build_safe_error("Dosya listesi alınamadı.")


@app.delete("/api/files/{source}")
def delete_file(source: str):
    source = sanitize_filename(source)
    
    try:
        collection = get_collection()
        # Find IDs for the source
        # ChromaDB get always returns IDs. include=[] minimizes data transfer.
        to_delete = collection.get(where={"source": source}, include=[])
        ids = to_delete.get("ids", [])
        
        if not ids:
            raise HTTPException(status_code=404, detail="Belirtilen kaynak bulunamadı.")
            
        collection.delete(ids=ids)
        return {"deleted": len(ids), "source": source}
    except HTTPException:
        raise
    except Exception:
        logger.exception("Silme hatası")
        raise build_safe_error("Dosya silinemedi.")


@app.post("/api/reindex")
def reindex():
    try:
        # Tüm koleksiyonu temizler
        client = get_chroma_client()
        client.delete_collection("documents")
        # yeniden oluştur
        client.get_or_create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )
        return {"status": "cleared"}
    except Exception:
        logger.exception("Sıfırlama hatası")
        raise build_safe_error("Koleksiyon sıfırlanamadı.")



@app.post("/api/embed")
async def embed_documents(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(status_code=400, detail="Dosya yüklenmedi.")
    if len(files) > MAX_FILE_COUNT:
        raise HTTPException(status_code=400, detail=f"En fazla {MAX_FILE_COUNT} dosya yüklenebilir.")

    logger.info(f"Yükleme başladı: {len(files)} dosya")
    docs: List[Document] = []
    errors = []
    
    for f in files:
        try:
            safe_name = sanitize_filename(f.filename)
            suffix = Path(safe_name).suffix.lower()
            if suffix not in ALLOWED_EXTENSIONS:
                errors.append(f"{safe_name}: Desteklenmeyen dosya formatı")
                continue
            if f.content_type and f.content_type not in ALLOWED_CONTENT_TYPES:
                errors.append(f"{safe_name}: Desteklenmeyen içerik tipi")
                continue

            logger.info(f"Dosya okunuyor: {safe_name}")
            data = await f.read()
            file_size_mb = len(data) / (1024 * 1024)
            logger.info(f"{safe_name} boyutu: {file_size_mb:.2f}MB")
            
            if file_size_mb > MAX_FILE_SIZE_MB:
                errors.append(f"{safe_name}: Dosya çok büyük ({file_size_mb:.1f}MB). Maksimum {MAX_FILE_SIZE_MB}MB.")
                continue
            
            content = ""
            if suffix == ".txt" or f.content_type == "text/plain":
                content = read_txt(io.BytesIO(data))
            elif suffix == ".pdf" or f.content_type == "application/pdf":
                try:
                    logger.info(f"PDF işleniyor: {safe_name}")
                    content = read_pdf(io.BytesIO(data))
                except Exception:
                    logger.exception("PDF okuma hatası: %s", safe_name)
                    errors.append(f"{safe_name}: PDF okunamadı")
                    continue
            elif suffix in {".docx", ".doc"} or f.content_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
                try:
                    logger.info(f"Word dosyası işleniyor: {safe_name}")
                    content = read_docx(io.BytesIO(data))
                except Exception:
                    logger.exception("Word okuma hatası: %s", safe_name)
                    errors.append(f"{safe_name}: Word dosyası okunamadı")
                    continue
            else:
                errors.append(f"{safe_name}: Desteklenmeyen dosya formatı")
                continue

            if content.strip():
                docs.append(Document(content=content, meta={"source": safe_name}))
                logger.info(f"{safe_name} başarıyla işlendi, içerik uzunluğu: {len(content)} karakter")
            else:
                errors.append(f"{safe_name}: Dosya boş veya içerik çıkarılamadı")
        except HTTPException:
            raise
        except Exception:
            logger.exception("İşleme hatası: %s", getattr(f, "filename", "bilinmeyen-dosya"))
            errors.append(f"{getattr(f, 'filename', 'Bilinmeyen dosya')}: İşleme hatası")
            continue

    if not docs:
        error_msg = "Geçerli doküman bulunamadı."
        if errors:
            error_msg += " Hatalar: " + "; ".join(errors[:3])
        raise HTTPException(status_code=400, detail=error_msg)

    try:
        logger.info(f"Vektörleştirme başlıyor: {len(docs)} doküman")
        collection = get_collection()
        embedder = load_embedder()
        added = embed_and_store(docs, collection, embedder)
        logger.info(f"Vektörleştirme tamamlandı: {added} parça eklendi")
        
        result = {"added_chunks": added, "files": [d.meta.get("source") for d in docs]}
        if errors:
            result["warnings"] = errors
        
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Vektörleştirme hatası")
        raise normalize_ollama_exception(exc=e, model_name=EMBEDDING_MODEL)


@app.post("/api/query")
def query_docs(body: QueryRequest):
    import time
    start_time = time.time()
    
    query = body.query.strip()
    if not query:
        raise HTTPException(status_code=400, detail="Soru boş olamaz.")
    top_k = body.top_k

    collection = get_collection()
    embedder = load_embedder()
    generator = load_generator()

    # Embedding time
    embed_start = time.time()
    query_embedding = embedder.run(text=query)["embedding"]
    embed_time = time.time() - embed_start
    
    # Retrieval time
    retrieval_start = time.time()
    results = collection.query(query_embeddings=[query_embedding], n_results=top_k)
    retrieval_time = time.time() - retrieval_start
    
    contexts = results.get("documents", [[]])[0]
    distances = results.get("distances", [[]])[0]
    metadatas = results.get("metadatas", [[]])[0]

    if not contexts:
        return {"answer": None, "contexts": [], "note": "Veritabanında sonuç yok"}

    style = (body.answer_style or "default").lower()
    if style == "short":
        style_instruction = "Çok kısa ve öz (1-2 cümle)."
    elif style == "bullet":
        style_instruction = "Madde madde yanıtla; 3-6 madde; kısa cümleler."
    elif style == "long":
        style_instruction = "Daha açıklayıcı, gerekirse birkaç cümle veya maddeler."
    else:
        style_instruction = "Standart, net yanıt."

    prompt = build_prompt(query, contexts, style_instruction)
    
    # Generation time
    gen_start = time.time()
    try:
        response = generator.run(prompt=prompt)
        reply = response["replies"][0]
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("LLM sorgu hatası")
        raise normalize_ollama_exception(exc=e, model_name=LLM_MODEL)
    gen_time = time.time() - gen_start
    
    total_time = time.time() - start_time
    
    # Calculate similarity scores from ChromaDB distances
    # ChromaDB with "cosine" space returns cosine distance: distance = 1 - cosine_similarity
    # Therefore: cosine_similarity = 1 - distance
    similarities = []
    if distances:
        for d in distances:
            # Cosine distance should be between 0 and 2
            # 0 = identical, 2 = opposite
            # Similarity = 1 - distance gives us 0-1 range
            similarity = max(0, min(1, 1 - d))  # Clamp to [0, 1]
            similarities.append(similarity)
    
    # Developer metadata
    metadata = {
        "total_time": round(total_time, 3),
        "embed_time": round(embed_time, 3),
        "retrieval_time": round(retrieval_time, 3),
        "generation_time": round(gen_time, 3),
        "token_count": len(reply.split()),
        "context_count": len(contexts),
        "avg_similarity": round(sum(similarities) / len(similarities), 3) if similarities else 0,
        "top_similarity": round(max(similarities), 3) if similarities else 0,
        "similarities": [round(s, 3) for s in similarities],
        "raw_distances": [round(d, 3) for d in distances] if distances else [],  # Debug için
        "sources": [m.get("source", "unknown") for m in metadatas] if metadatas else [],
        "prompt_length": len(prompt),
    }
    
    return {
        "answer": reply, 
        "contexts": contexts, 
        "prompt": prompt, 
        "style": style,
        "metadata": metadata
    }


@app.post("/api/query/stream")
async def query_docs_stream(body: QueryRequest):
    """Streaming version of query endpoint for real-time response"""
    import time
    
    start_time = time.time()
    
    query = body.query.strip()
    if not query:
        raise HTTPException(status_code=400, detail="Soru boş olamaz.")
    top_k = body.top_k

    collection = get_collection()
    embedder = load_embedder()

    embed_start = time.time()
    query_embedding = embedder.run(text=query)["embedding"]
    embed_time = time.time() - embed_start
    
    retrieval_start = time.time()
    results = collection.query(query_embeddings=[query_embedding], n_results=top_k)
    retrieval_time = time.time() - retrieval_start
    
    contexts = results.get("documents", [[]])[0]
    distances = results.get("distances", [[]])[0]
    metadatas = results.get("metadatas", [[]])[0]

    if not contexts:
        async def empty_stream():
            yield f"data: {json.dumps({'type': 'error', 'content': 'Veritabanında sonuç yok'})}\n\n"
        return StreamingResponse(empty_stream(), media_type="text/event-stream")

    style = (body.answer_style or "default").lower()
    if style == "short":
        style_instruction = "Çok kısa ve öz (1-2 cümle)."
    elif style == "bullet":
        style_instruction = "Madde madde yanıtla; 3-6 madde; kısa cümleler."
    elif style == "long":
        style_instruction = "Daha açıklayıcı, gerekirse birkaç cümle veya maddeler."
    else:
        style_instruction = "Standart, net yanıt."

    prompt = build_prompt(query, contexts, style_instruction)
    
    # Calculate similarities (same formula as non-streaming)
    similarities = []
    if distances:
        for d in distances:
            similarity = max(0, min(1, 1 - d))
            similarities.append(similarity)
    
    async def generate_stream():
        try:
            # İlk olarak context'leri gönder
            yield f"data: {json.dumps({'type': 'contexts', 'content': contexts})}\n\n"
            
            # Ollama'dan streaming response al
            gen_start = time.time()
            token_count = 0

            response = requests.post(
                f"{OLLAMA_URL}/api/generate",
                json={
                    "model": LLM_MODEL,
                    "prompt": prompt,
                    "stream": True,
                    "think": False,
                    "options": {
                        "num_predict": GENERATION_NUM_PREDICT,
                        "temperature": 0.7,
                    }
                },
                stream=True,
                timeout=(10, 120),
            )
            response.raise_for_status()
            
            for line in response.iter_lines():
                if line:
                    chunk = json.loads(line)
                    if "response" in chunk:
                        token = chunk["response"]
                        token_count += len(token.split())
                        yield f"data: {json.dumps({'type': 'token', 'content': token})}\n\n"
                    
                    if chunk.get("done", False):
                        gen_time = time.time() - gen_start
                        total_time = time.time() - start_time
                        
                        # Metadata gönder
                        metadata = {
                            "total_time": round(total_time, 3),
                            "embed_time": round(embed_time, 3),
                            "retrieval_time": round(retrieval_time, 3),
                            "generation_time": round(gen_time, 3),
                            "token_count": token_count,
                            "context_count": len(contexts),
                            "avg_similarity": round(sum(similarities) / len(similarities), 3) if similarities else 0,
                            "top_similarity": round(max(similarities), 3) if similarities else 0,
                            "similarities": [round(s, 3) for s in similarities],
                            "raw_distances": [round(d, 3) for d in distances] if distances else [],
                            "sources": [m.get("source", "unknown") for m in metadatas] if metadatas else [],
                            "prompt_length": len(prompt),
                        }
                        
                        yield f"data: {json.dumps({'type': 'metadata', 'content': metadata})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'content': ''})}\n\n"
                        break
                        
        except Exception as e:
            logger.exception("Streaming sorgu hatası")
            error = normalize_ollama_exception(exc=e, model_name=LLM_MODEL)
            yield f"data: {json.dumps({'type': 'error', 'content': error.detail})}\n\n"
    
    return StreamingResponse(generate_stream(), media_type="text/event-stream")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)